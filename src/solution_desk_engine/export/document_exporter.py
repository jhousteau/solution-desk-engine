"""Document export functionality for technical sales proposals."""

import subprocess  # nosec
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional

from rich.console import Console  # type: ignore

console = Console()


class ExportFormat(Enum):
    """Supported document export formats."""

    MARKDOWN = "md"
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"


class ExportResult:
    """Result of a document export operation."""

    def __init__(
        self,
        success: bool,
        output_path: Optional[Path] = None,
        error: Optional[str] = None,
    ):
        self.success = success
        self.output_path = output_path
        self.error = error


class DocumentExporter:
    """Handles export of technical sales documents to various formats."""

    def __init__(self, output_dir: Optional[Path] = None) -> None:
        """Initialize the document exporter.

        Args:
            output_dir: Directory for exported documents. Defaults to ./output/
        """
        self.output_dir = output_dir or Path("output")
        self.output_dir.mkdir(exist_ok=True)

    def export_document(
        self,
        source_path: Path,
        format_type: ExportFormat,
        output_name: Optional[str] = None,
    ) -> ExportResult:
        """Export a single document to the specified format.

        Args:
            source_path: Path to the source markdown file
            format_type: Target export format
            output_name: Optional custom output filename (without extension)

        Returns:
            ExportResult with success status and output path
        """
        if not source_path.exists():
            return ExportResult(
                success=False, error=f"Source file not found: {source_path}"
            )

        # Determine output filename
        if output_name:
            output_filename = f"{output_name}.{format_type.value}"
        else:
            output_filename = f"{source_path.stem}.{format_type.value}"

        output_path = self.output_dir / output_filename

        try:
            if format_type == ExportFormat.MARKDOWN:
                # Simple copy for markdown
                import shutil

                shutil.copy2(source_path, output_path)
                return ExportResult(success=True, output_path=output_path)

            elif format_type == ExportFormat.PDF:
                return self._export_to_pdf(source_path, output_path)

            elif format_type == ExportFormat.DOCX:
                return self._export_to_docx(source_path, output_path)

            elif format_type == ExportFormat.HTML:
                return self._export_to_html(source_path, output_path)

            else:
                return ExportResult(
                    success=False, error=f"Unsupported format: {format_type}"
                )

        except Exception as e:
            return ExportResult(success=False, error=f"Export failed: {str(e)}")

    def export_multiple_documents(
        self, source_files: List[Path], format_type: ExportFormat
    ) -> Dict[Path, ExportResult]:
        """Export multiple documents to the specified format.

        Args:
            source_files: List of source markdown files
            format_type: Target export format

        Returns:
            Dictionary mapping source paths to export results
        """
        results = {}

        for source_file in source_files:
            console.print(f"Exporting {source_file.name} to {format_type.value}...")
            result = self.export_document(source_file, format_type)
            results[source_file] = result

            if result.success:
                console.print(f"✓ Exported to: {result.output_path}", style="green")
            else:
                console.print(f"✗ Failed: {result.error}", style="red")

        return results

    def _export_to_pdf(self, source_path: Path, output_path: Path) -> ExportResult:
        """Export markdown to PDF using pandoc or weasyprint."""
        try:
            # Try pandoc first
            cmd = [
                "pandoc",
                str(source_path),
                "-o",
                str(output_path),
                "--pdf-engine=weasyprint",
                (
                    "--css=styles/professional.css"
                    if Path("styles/professional.css").exists()
                    else ""
                ),
            ]
            cmd = [arg for arg in cmd if arg]  # Remove empty args

            subprocess.run(cmd, check=True, capture_output=True, text=True)  # nosec
            return ExportResult(success=True, output_path=output_path)

        except (subprocess.CalledProcessError, FileNotFoundError):
            # Fallback to weasyprint directly
            try:
                import markdown  # type: ignore
                import weasyprint  # type: ignore

                # Convert markdown to HTML first
                with open(source_path, "r", encoding="utf-8") as f:
                    md_content = f.read()

                html_content = markdown.markdown(
                    md_content, extensions=["tables", "toc"]
                )

                # Add basic styling
                styled_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <style>
                        body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 2cm; }}
                        h1, h2, h3 {{ color: #2c3e50; }}
                        table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
                        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                        th {{ background-color: #f2f2f2; }}
                        code {{ background-color: #f4f4f4; padding: 2px 4px; }}
                        pre {{ background-color: #f4f4f4; padding: 1em; overflow-x: auto; }}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
                """

                # Convert to PDF
                weasyprint.HTML(string=styled_html).write_pdf(output_path)
                return ExportResult(success=True, output_path=output_path)

            except ImportError:
                return ExportResult(
                    success=False,
                    error="PDF export requires pandoc or weasyprint. Install with: pip install weasyprint",
                )
            except Exception as e:
                return ExportResult(success=False, error=f"PDF export failed: {str(e)}")

    def _export_to_docx(self, source_path: Path, output_path: Path) -> ExportResult:
        """Export markdown to DOCX using pandoc or python-docx."""
        try:
            # Try pandoc first
            cmd = [
                "pandoc",
                str(source_path),
                "-o",
                str(output_path),
                (
                    "--reference-doc=styles/reference.docx"
                    if Path("styles/reference.docx").exists()
                    else ""
                ),
            ]
            cmd = [arg for arg in cmd if arg]  # Remove empty args

            subprocess.run(cmd, check=True, capture_output=True, text=True)  # nosec
            return ExportResult(success=True, output_path=output_path)

        except (subprocess.CalledProcessError, FileNotFoundError):
            # Fallback to python-docx
            try:
                from docx import Document  # type: ignore

                # Convert markdown to HTML first
                with open(source_path, "r", encoding="utf-8") as f:
                    md_content = f.read()

                # Create a basic DOCX document
                doc = Document()

                # Simple markdown parsing for basic conversion
                lines = md_content.split("\n")
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    elif line.startswith("# "):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith("### "):
                        doc.add_heading(line[4:], level=3)
                    else:
                        doc.add_paragraph(line)

                doc.save(output_path)
                return ExportResult(success=True, output_path=output_path)

            except ImportError:
                return ExportResult(
                    success=False,
                    error="DOCX export requires pandoc or python-docx. Install with: pip install python-docx",
                )
            except Exception as e:
                return ExportResult(
                    success=False, error=f"DOCX export failed: {str(e)}"
                )

    def _export_to_html(self, source_path: Path, output_path: Path) -> ExportResult:
        """Export markdown to HTML."""
        try:
            import markdown

            with open(source_path, "r", encoding="utf-8") as f:
                md_content = f.read()

            html_content = markdown.markdown(
                md_content, extensions=["tables", "toc", "fenced_code", "codehilite"]
            )

            # Add basic HTML structure and styling
            styled_html = f"""
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="utf-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{source_path.stem}</title>
                <style>
                    body {{
                        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                        line-height: 1.6;
                        max-width: 1000px;
                        margin: 0 auto;
                        padding: 2rem;
                        color: #333;
                    }}
                    h1, h2, h3, h4, h5, h6 {{ color: #2c3e50; margin-top: 2rem; }}
                    h1 {{ border-bottom: 3px solid #3498db; padding-bottom: 0.5rem; }}
                    h2 {{ border-bottom: 2px solid #ecf0f1; padding-bottom: 0.3rem; }}
                    table {{
                        border-collapse: collapse;
                        width: 100%;
                        margin: 1.5rem 0;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                    }}
                    th, td {{
                        border: 1px solid #ddd;
                        padding: 12px;
                        text-align: left;
                    }}
                    th {{
                        background-color: #3498db;
                        color: white;
                        font-weight: bold;
                    }}
                    tr:nth-child(even) {{ background-color: #f8f9fa; }}
                    code {{
                        background-color: #f4f4f4;
                        padding: 2px 4px;
                        border-radius: 3px;
                        font-family: 'Monaco', 'Menlo', monospace;
                    }}
                    pre {{
                        background-color: #f8f9fa;
                        padding: 1rem;
                        border-radius: 5px;
                        overflow-x: auto;
                        border-left: 4px solid #3498db;
                    }}
                    blockquote {{
                        border-left: 4px solid #3498db;
                        margin: 1rem 0;
                        padding: 0.5rem 1rem;
                        background-color: #f8f9fa;
                    }}
                    .toc {{
                        background-color: #ecf0f1;
                        padding: 1rem;
                        border-radius: 5px;
                        margin: 2rem 0;
                    }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(styled_html)

            return ExportResult(success=True, output_path=output_path)

        except ImportError:
            return ExportResult(
                success=False,
                error="HTML export requires markdown. Install with: pip install markdown",
            )
        except Exception as e:
            return ExportResult(success=False, error=f"HTML export failed: {str(e)}")

    def get_export_summary(self, results: Dict[Path, ExportResult]) -> Dict[str, Any]:
        """Generate a summary of export operations.

        Args:
            results: Dictionary of export results

        Returns:
            Summary dictionary with counts and details
        """
        successful = sum(1 for result in results.values() if result.success)
        failed = len(results) - successful

        return {
            "total_files": len(results),
            "successful": successful,
            "failed": failed,
            "success_rate": (successful / len(results) * 100) if results else 0,
            "output_files": [
                str(result.output_path)
                for result in results.values()
                if result.success and result.output_path
            ],
            "errors": [
                result.error
                for result in results.values()
                if not result.success and result.error
            ],
        }
