#!/usr/bin/env python3
"""
SOW Template Processor for Google DAF/PSF Documents
Specialized processor for creating Statement of Work documents that match Google's
professional formatting standards and DAF/PSF compliance requirements.
"""

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    from docx.table import _Cell
except ImportError:
    print("Error: python-docx not installed")
    exit(1)

try:
    from branding_manager import BrandingManager
except ImportError:
    print("Error: branding_manager.py not found")
    exit(1)


class SOWTemplateProcessor:
    """Specialized processor for Google DAF/PSF SOW templates."""

    def __init__(self, branding_dir: Path):
        self.branding = BrandingManager(branding_dir)
        self.logger = logging.getLogger(__name__)

        # Google SOW specific colors (complement to Capgemini brand colors)
        self.google_colors = {
            "google_blue": RGBColor(66, 133, 244),  # #4285f4 - Google Blue
            "google_red": RGBColor(234, 67, 53),  # #ea4335 - Google Red
            "google_yellow": RGBColor(251, 188, 5),  # #fbbc05 - Google Yellow
            "google_green": RGBColor(52, 168, 83),  # #34a853 - Google Green
            "doc_header": RGBColor(68, 68, 68),  # #444444 - Document headers
            "table_border": RGBColor(204, 204, 204),  # #cccccc - Table borders
        }

    def create_sow_cover_page(self, doc: Document, metadata: Dict) -> None:
        """Create professional SOW cover page matching Google standards."""
        # Add partner and customer logos side by side
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add Capgemini logo
        if self.branding.logo_color_png.exists():
            try:
                logo_para.add_run().add_picture(
                    str(self.branding.logo_color_png), width=Inches(2)
                )
                logo_para.add_run("     ")  # Spacing between logos
                # Placeholder for customer logo
                logo_para.add_run("[Customer Logo]")
            except Exception as e:
                self.logger.warning(f"Could not add logo: {e}")

        # Add spacing
        doc.add_paragraph()
        doc.add_paragraph()

        # Main title
        title_para = doc.add_heading("Google Cloud Workload Implementation", level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.style.font.size = Pt(24)
        title_para.style.font.color.rgb = self.branding.brand_colors_rgb["primary_blue"]
        title_para.style.font.bold = True

        # Statement of Work subtitle
        sow_para = doc.add_heading("Statement of Work", level=2)
        sow_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sow_para.style.font.size = Pt(20)
        sow_para.style.font.color.rgb = self.branding.brand_colors_rgb["dark_blue"]

        # Add spacing
        for _ in range(3):
            doc.add_paragraph()

        # Partner and Customer names
        partner_para = doc.add_heading("[PARTNER NAME]", level=2)
        partner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        partner_para.style.font.size = Pt(18)
        partner_para.style.font.color.rgb = self.branding.brand_colors_rgb["text_gray"]

        customer_para = doc.add_heading("[CUSTOMER NAME]", level=2)
        customer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        customer_para.style.font.size = Pt(18)
        customer_para.style.font.color.rgb = self.branding.brand_colors_rgb["text_gray"]

        # Add spacing
        for _ in range(4):
            doc.add_paragraph()

        # Document metadata
        date_para = doc.add_paragraph(
            f'Date: {metadata.get("date", datetime.now().strftime("%Y-%m-%d"))}'
        )
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.style.font.size = Pt(12)

        author_para = doc.add_paragraph(
            f'Author: {metadata.get("author", "Capgemini")}'
        )
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_para.style.font.size = Pt(12)

        # Confidentiality notice
        for _ in range(3):
            doc.add_paragraph()

        conf_para = doc.add_paragraph("Confidential - Capgemini & Customer")
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_para.style.font.size = Pt(10)
        conf_para.style.font.color.rgb = self.branding.brand_colors_rgb["text_gray"]
        conf_para.style.font.italic = True

        # Page break
        doc.add_page_break()

    def setup_sow_styles(self, doc: Document) -> None:
        """Setup specialized styles for SOW documents."""
        # Base branding styles
        self.branding.setup_docx_styles(doc)

        # SOW-specific table style
        if "SOW Table" not in doc.styles:
            table_style = doc.styles.add_style("SOW Table", WD_STYLE_TYPE.TABLE)
            table_style.base_style = doc.styles["Table Grid"]

        # Gantt chart style
        if "Gantt Chart" not in doc.styles:
            gantt_style = doc.styles.add_style("Gantt Chart", WD_STYLE_TYPE.TABLE)

        # Section header style
        if "SOW Section" not in doc.styles:
            section_style = doc.styles.add_style("SOW Section", WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = self.branding.brand_fonts["primary"]
            section_style.font.size = Pt(14)
            section_style.font.color.rgb = self.google_colors["doc_header"]
            section_style.font.bold = True
            section_style.paragraph_format.space_before = Pt(18)
            section_style.paragraph_format.space_after = Pt(6)

    def format_sow_table(
        self,
        doc: Document,
        table_data: List[List[str]],
        has_header: bool = True,
        table_type: str = "standard",
    ) -> None:
        """Format tables with SOW-appropriate styling."""
        if not table_data:
            return

        # Create table
        num_rows = len(table_data)
        num_cols = len(table_data[0]) if table_data else 0

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Apply table-specific formatting
        for row_idx, row_data in enumerate(table_data):
            table_row = table.rows[row_idx]

            for col_idx, cell_data in enumerate(row_data):
                cell = table_row.cells[col_idx]
                cell.text = str(cell_data)

                # Header row formatting
                if has_header and row_idx == 0:
                    self._format_table_header_cell(cell)
                # Gantt chart formatting
                elif table_type == "gantt" and col_idx > 0:
                    self._format_gantt_cell(cell, cell_data)
                # Standard cell formatting
                else:
                    self._format_standard_table_cell(cell, row_idx)

    def _format_table_header_cell(self, cell: _Cell) -> None:
        """Format table header cells with SOW styling."""
        # Set background color
        self._set_cell_background(cell, self.branding.brand_colors_rgb["primary_blue"])

        # Format text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = self.branding.brand_colors_rgb["white"]
                run.font.bold = True
                run.font.size = Pt(10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _format_gantt_cell(self, cell: _Cell, content: str) -> None:
        """Format Gantt chart cells with visual indicators."""
        # Detect Gantt chart symbols
        if "■" in content or "█" in content:
            # Filled cell - active period
            self._set_cell_background(
                cell, self.branding.brand_colors_rgb["primary_blue"]
            )
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = self.branding.brand_colors_rgb["white"]
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif content.strip():
            # Text cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = self.branding.brand_colors_rgb["text_gray"]

    def _format_standard_table_cell(self, cell: _Cell, row_idx: int) -> None:
        """Format standard table cells with alternating colors."""
        # Alternating row colors
        if row_idx % 2 == 1:
            self._set_cell_background(
                cell, self.branding.brand_colors_rgb["light_gray"]
            )

        # Standard text formatting
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = self.branding.brand_colors_rgb["text_gray"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _set_cell_background(self, cell: _Cell, color: RGBColor) -> None:
        """Set cell background color."""
        try:
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), f"{color.r:02x}{color.g:02x}{color.b:02x}")
            cell._element.get_or_add_tcPr().append(shading_elm)
        except Exception as e:
            self.logger.warning(f"Could not set cell background: {e}")

    def create_table_of_contents(self, doc: Document, content: str) -> None:
        """Create professional table of contents from markdown headings."""
        toc_title = doc.add_heading("Contents", level=1)
        toc_title.style.font.color.rgb = self.branding.brand_colors_rgb["primary_blue"]
        toc_title.style.font.size = Pt(16)

        # Extract headings from content
        headings = self._extract_headings(content)

        for heading_level, heading_text, page_num in headings:
            toc_line = doc.add_paragraph()

            # Add indentation based on heading level
            indent = (heading_level - 1) * 0.25
            toc_line.paragraph_format.left_indent = Inches(indent)

            # Add heading text
            text_run = toc_line.add_run(heading_text)
            text_run.font.size = Pt(11)
            text_run.font.color.rgb = self.branding.brand_colors_rgb["text_gray"]

            # Add dots and page number
            dots_run = toc_line.add_run(
                " " + "." * (60 - len(heading_text) - len(str(page_num)))
            )
            dots_run.font.size = Pt(11)
            dots_run.font.color.rgb = self.branding.brand_colors_rgb["medium_gray"]

            page_run = toc_line.add_run(f" {page_num}")
            page_run.font.size = Pt(11)
            page_run.font.color.rgb = self.branding.brand_colors_rgb["text_gray"]
            page_run.font.bold = True

        doc.add_page_break()

    def _extract_headings(self, content: str) -> List[Tuple[int, str, int]]:
        """Extract headings from markdown content."""
        headings = []
        page_counter = 5  # Start after cover page and TOC

        for line in content.split("\n"):
            if line.strip().startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                text = line.lstrip("#").strip()

                if text and not text.startswith("[DELETE]"):
                    headings.append((level, text, page_counter))
                    page_counter += 1

        return headings

    def add_signature_block(self, doc: Document) -> None:
        """Add professional signature block."""
        doc.add_paragraph()

        # Signature instructions
        sig_intro = doc.add_paragraph(
            "By signing below, [CUSTOMER NAME] confirms that [PARTNER NAME] has "
            "completed the deliverables as stated in this SOW."
        )
        sig_intro.style.font.size = Pt(11)

        doc.add_paragraph()

        # Create signature table
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.style = "Table Grid"

        # Customer signature
        customer_cell = sig_table.rows[0].cells[0]
        customer_cell.text = "Customer\n\nSignature: _________________________\nName: _________________________\nTitle: _________________________\nDate: _________________________"

        # Partner signature
        partner_cell = sig_table.rows[0].cells[1]
        partner_cell.text = "Partner\n\nSignature: _________________________\nName: _________________________\nTitle: _________________________\nDate: _________________________"

        # Format signature cells
        for row in sig_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                for paragraph in cell.paragraphs:
                    paragraph.style.font.size = Pt(10)
                    paragraph.style.font.color.rgb = self.branding.brand_colors_rgb[
                        "text_gray"
                    ]

    def process_sow_content(self, content: str) -> str:
        """Preprocess SOW content for optimal formatting."""
        # Remove template instructions
        content = re.sub(r"\[DELETE\].*?\n", "", content, flags=re.MULTILINE)
        content = re.sub(r"-ToBeDeleted-.*?-ToBeDeleted-", "", content, flags=re.DOTALL)
        content = re.sub(
            r"/CanBeRemoved/.*?/CanBeRemoved/", "", content, flags=re.DOTALL
        )

        # Enhance table formatting markers
        content = re.sub(r"\| Weeks \|", "| **Timeline** |", content)
        content = re.sub(r"\| Role \|", "| **Role** |", content)
        content = re.sub(
            r"\| Workstreams Role Efforts \|", "| **Workstreams** |", content
        )

        # Add page breaks before major sections
        major_sections = [
            "## 1. Executive Summary",
            "## 2. ",
            "## 3. Activities and Deliverables",
            "## 4. Prerequisites",
            "## 5. Roles & Responsibilities",
            "## 6. Financials",
            "## 7. Acceptance Criteria",
            "## 8. Appendix",
        ]

        for section in major_sections:
            content = content.replace(
                section, f'<div class="page-break"></div>\n\n{section}'
            )

        return content

    def enhance_sow_metadata(self, metadata: Dict, file_path: Path) -> Dict:
        """Enhance metadata with SOW-specific information."""
        sow_metadata = self.branding.get_brand_metadata(
            metadata.get("title", "Google Cloud Workload Implementation SOW")
        )

        # SOW-specific fields
        sow_metadata.update(
            {
                "document_type": "Statement of Work",
                "template_type": "Google DAF/PSF SOW Template Y25",
                "compliance": "DAF/PSF Requirements",
                "version": "2025.1",
                "classification": "Confidential",
                "approval_required": "Yes",
                "funding_type": "DAF/PSF",
                "region": "NORTHAM",  # Default, should be customized
            }
        )

        sow_metadata.update(metadata)
        return sow_metadata
