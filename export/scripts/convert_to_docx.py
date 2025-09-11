#!/usr/bin/env python3
"""
DOCX Conversion Script for [CLIENT] F&I Documentation
Converts markdown files to professionally styled Word documents using python-docx library.
Follows Capgemini brand guidelines from sow-word-conversion-guide.md
"""

import argparse
import json
import logging
import re
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("Error: python-docx not installed. Run: poetry add python-docx")
    sys.exit(1)

try:
    from branding_manager import BrandingManager
    from sow_template_processor import SOWTemplateProcessor
except ImportError:
    print(
        "Error: branding_manager.py or sow_template_processor.py not found. Ensure they exist in the same directory."
    )
    sys.exit(1)

try:
    import mistune
except ImportError:
    print("Error: mistune not installed. Run: poetry add mistune")
    sys.exit(1)

# Configuration
OPPORTUNITY_DIR = Path("../../opportunity")
OUTPUT_DIR = Path("../docx")
BRANDING_DIR = Path("../../branding")
METADATA_DIR = Path("../_metadata")

# Capgemini Brand Colors
COLORS = {
    "primary_blue": RGBColor(0, 112, 173),  # #0070AD
    "secondary_teal": RGBColor(18, 171, 219),  # #12ABDB
    "dark_blue": RGBColor(0, 51, 161),  # #0033A1
    "text_gray": RGBColor(51, 51, 51),  # #333333
    "light_gray": RGBColor(245, 245, 245),  # #F5F5F5
    "medium_gray": RGBColor(128, 128, 128),  # #808080
    "white": RGBColor(255, 255, 255),  # #FFFFFF
}

# Phase configuration
PHASES = {
    "0-source": {"name": "Source Materials", "color": COLORS["dark_blue"]},
    "1-research": {"name": "Research & Discovery", "color": COLORS["primary_blue"]},
    "2-requirements": {
        "name": "Requirements Analysis",
        "color": COLORS["primary_blue"],
    },
    "3-analysis": {"name": "Stakeholder Analysis", "color": COLORS["secondary_teal"]},
    "4-business-case": {"name": "Business Case", "color": COLORS["dark_blue"]},
    "5-architecture": {"name": "Architecture Design", "color": COLORS["primary_blue"]},
    "6-solution-design": {"name": "Solution Design", "color": COLORS["secondary_teal"]},
    "7-implementation-plan": {
        "name": "Implementation Plan",
        "color": COLORS["dark_blue"],
    },
    "8-proposal": {"name": "Proposal Materials", "color": COLORS["primary_blue"]},
    "9-contract": {"name": "Contract Documents", "color": COLORS["dark_blue"]},
    "10-audit": {"name": "Audit & Validation", "color": COLORS["text_gray"]},
}


class DocumentProcessor:
    """Processes markdown documents and converts them to DOCX."""

    def __init__(self, output_dir: Path, verbose: bool = False):
        self.output_dir = output_dir
        self.verbose = verbose
        self.conversion_log = []
        self.failed_conversions = []

        # Setup logging
        logging.basicConfig(
            level=logging.INFO if verbose else logging.WARNING,
            format="%(asctime)s - %(levelname)s - %(message)s",
        )
        self.logger = logging.getLogger(__name__)

        # Initialize branding manager
        self.branding = BrandingManager(BRANDING_DIR)

        # Initialize SOW template processor
        self.sow_processor = SOWTemplateProcessor(BRANDING_DIR)

        # Legacy logo path for backward compatibility
        self.capgemini_logo = self.branding.logo_color_png

    def setup_document_styles(self, doc: Document) -> None:
        """Setup custom styles for the document using branding manager."""
        self.branding.setup_docx_styles(doc)

    def add_cover_page(self, doc: Document, metadata: Dict) -> None:
        """Add a branded cover page to the document."""
        # Get branded cover content
        title = metadata.get("title", "F&I AI Assistant - Proof of Value (POV)")
        cover_content = self.branding.create_brand_cover_page_content(
            title=title, date=metadata.get("date", datetime.now().strftime("%Y-%m-%d"))
        )

        # Add logo if exists
        if self.branding.logo_color_png.exists():
            try:
                doc.add_picture(str(self.branding.logo_color_png), width=Inches(2))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            except Exception as e:
                self.logger.warning(f"Could not add logo: {e}")

        # Add spacing
        doc.add_paragraph()
        doc.add_paragraph()

        # Add title with brand colors
        title_heading = doc.add_heading("Statement of Work", level=1)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_heading.style.font.size = Pt(self.branding.typography_scale["h1"])
        title_heading.style.font.color.rgb = self.branding.brand_colors_rgb[
            "primary_blue"
        ]

        # Add subtitle with brand colors
        subtitle = doc.add_heading("[CLIENT] Automotive Group", level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.style.font.size = Pt(self.branding.typography_scale["h2"])
        subtitle.style.font.color.rgb = self.branding.brand_colors_rgb["dark_blue"]

        # Add project title
        project = doc.add_heading(title, level=2)
        project.alignment = WD_ALIGN_PARAGRAPH.CENTER
        project.style.font.size = Pt(18)
        project.style.font.color.rgb = self.branding.brand_colors_rgb["dark_blue"]

        doc.add_paragraph()

        # Add DAF implementation note
        daf = doc.add_paragraph(
            "Google Cloud Deal Acceleration Fund (DAF) Implementation"
        )
        daf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        daf.style.font.size = Pt(14)

        # Add spacing before date/version
        for _ in range(5):
            doc.add_paragraph()

        # Add date and version with brand metadata
        brand_metadata = self.branding.get_brand_metadata(title)
        date_para = doc.add_paragraph(f'Date: {brand_metadata["date"]}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        version_para = doc.add_paragraph(f'Version: {brand_metadata["version"]}')
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add confidentiality notice
        conf_para = doc.add_paragraph(
            f'{brand_metadata["confidentiality"]} - {brand_metadata["company"]}'
        )
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_para.style.font.size = Pt(self.branding.typography_scale["small"])
        conf_para.style.font.color.rgb = self.branding.brand_colors_rgb["text_gray"]

        # Add page break
        doc.add_page_break()

    def add_header_footer(self, doc: Document) -> None:
        """Add branded headers and footers to the document."""
        self.branding.apply_docx_header_footer(doc)

    def extract_metadata(self, content: str, file_path: Path) -> Dict:
        """Extract metadata from markdown frontmatter."""
        metadata = {
            "title": "",
            "author": "Capgemini",
            "date": datetime.now().strftime("%Y-%m-%d"),
            "phase": self.get_phase_from_path(file_path),
            "document_type": "documentation",
        }

        # Extract YAML frontmatter
        if content.startswith("---"):
            try:
                end_marker = content.find("---", 3)
                if end_marker != -1:
                    frontmatter = content[3:end_marker].strip()
                    for line in frontmatter.split("\n"):
                        if ":" in line:
                            key, value = line.split(":", 1)
                            key = key.strip()
                            value = value.strip().strip("\"'")
                            if key in metadata:
                                metadata[key] = value
            except Exception as e:
                self.logger.warning(f"Error parsing frontmatter: {e}")

        # Extract title from first H1 if not in frontmatter
        if not metadata["title"]:
            h1_match = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
            if h1_match:
                metadata["title"] = h1_match.group(1).strip()
            else:
                # Use filename as title
                metadata["title"] = file_path.stem.replace("-", " ").title()

        return metadata

    def _is_sow_template(self, content: str, file_path: Path) -> bool:
        """Detect if this is a SOW template file."""
        # Check file path
        if (
            "sow" in str(file_path).lower()
            or "statement-of-work" in str(file_path).lower()
        ):
            return True

        # Check content for SOW indicators
        sow_indicators = [
            "Statement of Work",
            "DAF",
            "PSF",
            "Google Cloud Workload Implementation",
            "Activities and Deliverables (Scope)",
            "Prerequisites and Customer Cooperation",
            "Roles & Responsibilities",
            "Acceptance Criteria",
            "[PARTNER NAME]",
            "[CUSTOMER NAME]",
        ]

        content_lower = content.lower()
        indicator_count = sum(
            1 for indicator in sow_indicators if indicator.lower() in content_lower
        )

        # If 3+ indicators present, treat as SOW template
        return indicator_count >= 3

    def get_phase_from_path(self, file_path: Path) -> Optional[str]:
        """Extract phase from file path."""
        path_parts = file_path.parts
        for part in path_parts:
            if part in PHASES:
                return part
        return None

    def preprocess_content(self, content: str) -> str:
        """Preprocess markdown content for better conversion."""
        # Remove frontmatter if present
        if content.startswith("---"):
            end_marker = content.find("---", 3)
            if end_marker != -1:
                content = content[end_marker + 3 :].strip()

        # Handle special markdown elements
        # Convert checkbox lists to regular lists
        content = re.sub(r"^\s*-\s*\[\s*\]\s*", "- [ ] ", content, flags=re.MULTILINE)
        content = re.sub(r"^\s*-\s*\[x\]\s*", "- [X] ", content, flags=re.MULTILINE)

        return content

    def convert_markdown_to_docx(self, doc: Document, content: str) -> None:
        """Convert markdown content to docx using mistune v3."""
        # Convert markdown to HTML first
        markdown = mistune.create_markdown(
            plugins=["table", "strikethrough", "footnotes", "task_lists"]
        )
        html_content = markdown(content)

        # For now, let's use a simpler approach - process markdown line by line
        self._process_markdown_lines(doc, content)

    def _process_markdown_lines(self, doc: Document, content: str) -> None:
        """Process markdown content line by line."""
        lines = content.split("\n")
        in_code_block = False
        in_table = False
        table_lines = []
        list_type = None

        for line in lines:
            # Handle code blocks
            if line.strip().startswith("```"):
                if in_code_block:
                    in_code_block = False
                else:
                    in_code_block = True
                    continue

            if in_code_block:
                p = doc.add_paragraph()
                if "Code" in doc.styles:
                    p.style = "Code"
                run = p.add_run(line)
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                continue

            # Handle tables
            if "|" in line and not in_table:
                in_table = True
                table_lines = [line]
            elif in_table and "|" in line:
                table_lines.append(line)
            elif in_table and "|" not in line and line.strip():
                # End of table
                self._process_simple_table(doc, table_lines)
                table_lines = []
                in_table = False
                # Process current line normally

            if in_table:
                continue

            # Handle headings
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                text = line.lstrip("#").strip()
                if text:
                    p = doc.add_heading(text, level=min(level, 3))
                    if level == 1:
                        p.style.font.color.rgb = COLORS["primary_blue"]
                        p.style.font.size = Pt(16)
                    elif level == 2:
                        p.style.font.color.rgb = COLORS["dark_blue"]
                        p.style.font.size = Pt(14)
                    else:
                        p.style.font.color.rgb = COLORS["text_gray"]
                        p.style.font.size = Pt(12)

            # Handle lists
            elif line.strip().startswith(("- ", "* ", "+ ")):
                p = doc.add_paragraph(style="List Bullet")
                self._add_formatted_text_simple(p, line.strip()[2:])
            elif re.match(r"^\s*\d+\.\s", line):
                p = doc.add_paragraph(style="List Number")
                text = re.sub(r"^\s*\d+\.\s", "", line)
                self._add_formatted_text_simple(p, text)

            # Handle blockquotes
            elif line.strip().startswith(">"):
                text = line.strip().lstrip(">").strip()
                if text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    if "Quote" in doc.styles:
                        p.style = "Quote"
                    self._add_formatted_text_simple(p, text)

            # Handle horizontal rules
            elif line.strip() in ["---", "***", "___"]:
                p = doc.add_paragraph("_" * 50)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Regular paragraphs
            elif line.strip():
                p = doc.add_paragraph()
                self._add_formatted_text_simple(p, line)

        # Handle any remaining table
        if table_lines:
            self._process_simple_table(doc, table_lines)

    def _add_formatted_text_simple(self, paragraph, text: str) -> None:
        """Add text with simple markdown formatting."""
        # Pattern to match bold, italic, and code
        pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif (
                part.startswith("*")
                and part.endswith("*")
                and not part.startswith("**")
            ):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Courier New"
            else:
                paragraph.add_run(part)

    def _process_simple_table(self, doc: Document, lines: list) -> None:
        """Process a simple markdown table."""
        if len(lines) < 3:  # Need header, separator, and at least one row
            return

        # Parse header
        headers = [cell.strip() for cell in lines[0].split("|") if cell.strip()]

        # Skip separator line

        # Parse rows
        rows = []
        for line in lines[2:]:
            if "|" in line:
                row = [cell.strip() for cell in line.split("|") if cell.strip()]
                if len(row) == len(headers):
                    rows.append(row)

        if not rows:
            return

        # Create table
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            self._style_header_cell(hdr_cells[i])

        # Add rows
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_text in enumerate(row_data):
                row_cells[col_idx].text = cell_text
                # Alternate row coloring
                if row_idx % 2 == 1:
                    self._set_cell_background(row_cells[col_idx], COLORS["light_gray"])

    def _process_tokens(
        self, doc: Document, tokens: list, in_list: bool = False
    ) -> None:
        """Process mistune tokens and add to document."""
        for token in tokens:
            try:
                token_type = token.get("type", "")
            except (TypeError, AttributeError) as e:
                self.logger.error(f"Error accessing token type: {e}, token: {token}")
                continue

            if token_type == "heading":
                level = token.get("attrs", {}).get("level", 1)
                text = self._extract_text(token.get("children", []))
                p = doc.add_heading(text, level=min(level, 3))
                if level == 1:
                    p.style.font.color.rgb = COLORS["primary_blue"]
                    p.style.font.size = Pt(16)
                elif level == 2:
                    p.style.font.color.rgb = COLORS["dark_blue"]
                    p.style.font.size = Pt(14)
                else:
                    p.style.font.color.rgb = COLORS["text_gray"]
                    p.style.font.size = Pt(12)

            elif token_type == "paragraph":
                if not in_list:
                    p = doc.add_paragraph()
                    self._add_inline_content(p, token.get("children", []))

            elif token_type == "list":
                ordered = token.get("attrs", {}).get("ordered", False)
                for item in token.get("children", []):
                    if item.get("type") == "list_item":
                        style = "List Number" if ordered else "List Bullet"
                        p = doc.add_paragraph(style=style)
                        if item.get("children"):
                            first_child = (
                                item["children"][0] if item["children"] else {}
                            )
                            if first_child.get("type") == "paragraph":
                                self._add_inline_content(
                                    p, first_child.get("children", [])
                                )
                            else:
                                self._add_inline_content(p, item["children"])

            elif token_type == "block_quote":
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                if "Quote" in doc.styles:
                    p.style = "Quote"
                self._process_tokens(doc, token.get("children", []))

            elif token_type == "block_code":
                p = doc.add_paragraph()
                if "Code" in doc.styles:
                    p.style = "Code"
                code_text = token.get("raw", "").strip("`").strip()
                run = p.add_run(code_text)
                run.font.name = "Courier New"
                run.font.size = Pt(10)

            elif token_type == "table":
                self._add_table(doc, token)

            elif token_type == "thematic_break":
                p = doc.add_paragraph("_" * 50)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _extract_text(self, children: list) -> str:
        """Extract plain text from token children."""
        text = ""
        for child in children:
            if isinstance(child, str):
                text += child
                continue
            child_type = child.get("type", "")
            if child_type == "text":
                text += child.get("raw", "")
            elif child_type == "strong":
                text += self._extract_text(child.get("children", []))
            elif child_type == "emphasis":
                text += self._extract_text(child.get("children", []))
            elif child_type == "code_span":
                text += child.get("raw", "").strip("`")
            elif child_type == "link":
                text += self._extract_text(child.get("children", []))
        return text

    def _add_inline_content(self, paragraph, children: list) -> None:
        """Add inline content to a paragraph."""
        for child in children:
            if isinstance(child, str):
                paragraph.add_run(child)
                continue
            child_type = child.get("type", "")
            if child_type == "text":
                paragraph.add_run(child.get("raw", ""))
            elif child_type == "strong":
                text = self._extract_text(child.get("children", []))
                run = paragraph.add_run(text)
                run.bold = True
            elif child_type == "emphasis":
                text = self._extract_text(child.get("children", []))
                run = paragraph.add_run(text)
                run.italic = True
            elif child_type == "code_span":
                text = child.get("raw", "").strip("`")
                run = paragraph.add_run(text)
                run.font.name = "Courier New"
            elif child_type == "link":
                text = self._extract_text(child.get("children", []))
                paragraph.add_run(text)  # Simple text for now
            elif child_type == "linebreak":
                paragraph.add_run("\n")

    def _add_table(self, doc: Document, token: dict) -> None:
        """Add a table to the document."""
        # Extract table data
        headers = []
        rows = []

        for child in token.get("children", []):
            child_type = child.get("type", "")
            if child_type == "table_head":
                for row in child.get("children", []):
                    if row.get("type") == "table_row":
                        headers = [
                            self._extract_text(cell.get("children", []))
                            for cell in row.get("children", [])
                        ]
            elif child_type == "table_body":
                for row in child.get("children", []):
                    if row.get("type") == "table_row":
                        row_data = [
                            self._extract_text(cell.get("children", []))
                            for cell in row.get("children", [])
                        ]
                        rows.append(row_data)

        if headers:
            # Create table
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Add headers
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                self._style_header_cell(hdr_cells[i])

            # Add rows
            for row_idx, row_data in enumerate(rows):
                row_cells = table.rows[row_idx + 1].cells
                for col_idx, cell_text in enumerate(row_data):
                    row_cells[col_idx].text = cell_text
                    # Alternate row coloring
                    if row_idx % 2 == 1:
                        self._set_cell_background(
                            row_cells[col_idx], COLORS["light_gray"]
                        )

    def _style_header_cell(self, cell):
        """Style table header cell."""
        self._set_cell_background(cell, COLORS["primary_blue"])
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = COLORS["white"]
                run.font.bold = True

    def _set_cell_background(self, cell, color):
        """Set cell background color."""
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), f"{color[0]:02x}{color[1]:02x}{color[2]:02x}")
        cell._element.get_or_add_tcPr().append(shading_elm)

    def convert_file(self, input_path: Path, output_path: Path) -> bool:
        """Convert a single markdown file to DOCX."""
        try:
            # Read markdown content
            with open(input_path, "r", encoding="utf-8") as f:
                content = f.read()

            # Extract metadata
            metadata = self.extract_metadata(content, input_path)

            # Detect if this is a SOW template
            is_sow_template = self._is_sow_template(content, input_path)

            # Create new document
            doc = Document()

            # Setup styles (SOW-specific if needed)
            if is_sow_template:
                self.sow_processor.setup_sow_styles(doc)
            else:
                self.setup_document_styles(doc)

            # Add cover page
            phase = self.get_phase_from_path(input_path)
            if is_sow_template:
                # Enhanced SOW cover page
                sow_metadata = self.sow_processor.enhance_sow_metadata(
                    metadata, input_path
                )
                self.sow_processor.create_sow_cover_page(doc, sow_metadata)
                # Add table of contents for SOW
                self.sow_processor.create_table_of_contents(doc, content)
            elif phase == "9-contract":
                self.add_cover_page(doc, metadata)

            # Add header/footer
            self.add_header_footer(doc)

            # Preprocess content
            if is_sow_template:
                processed_content = self.sow_processor.process_sow_content(content)
            else:
                processed_content = self.preprocess_content(content)

            # Convert markdown to docx
            self.convert_markdown_to_docx(doc, processed_content)

            # Add signature block for SOW templates
            if is_sow_template:
                self.sow_processor.add_signature_block(doc)

            # Ensure output directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # Save document
            doc.save(str(output_path))

            # Log success
            self.conversion_log.append(
                {
                    "input": str(input_path),
                    "output": str(output_path),
                    "status": "success",
                    "metadata": metadata,
                    "timestamp": datetime.now().isoformat(),
                }
            )

            if self.verbose:
                self.logger.info(f"Converted: {input_path} -> {output_path}")

            return True

        except Exception as e:
            self.logger.error(f"Error converting {input_path}: {e}")
            self.failed_conversions.append(
                {
                    "input": str(input_path),
                    "error": str(e),
                    "timestamp": datetime.now().isoformat(),
                }
            )
            return False

    def find_markdown_files(self, directory: Path) -> List[Path]:
        """Find all markdown files in directory."""
        if not directory.exists():
            self.logger.warning(f"Directory not found: {directory}")
            return []

        md_files = []
        for file_path in directory.rglob("*.md"):
            if file_path.is_file():
                # Skip certain files
                if file_path.name in ["README.md", "CLAUDE.md"]:
                    continue
                md_files.append(file_path)

        return sorted(md_files)

    def create_output_structure(self, input_dir: Path, output_dir: Path) -> None:
        """Create output directory structure mirroring input."""
        for phase in PHASES:
            phase_dir = output_dir / phase
            phase_dir.mkdir(parents=True, exist_ok=True)

        # Create metadata directory
        METADATA_DIR.mkdir(parents=True, exist_ok=True)

    def process_directory(self, input_dir: Path, max_workers: int = 4) -> Dict:
        """Process all markdown files in directory."""
        # Find all markdown files
        md_files = self.find_markdown_files(input_dir)

        if not md_files:
            self.logger.warning(f"No markdown files found in {input_dir}")
            return {"total": 0, "successful": 0, "failed": 0}

        # Create output structure
        self.create_output_structure(input_dir, self.output_dir)

        # Process files
        successful = 0
        failed = 0

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all conversion tasks
            future_to_file = {}
            for md_file in md_files:
                # Calculate relative path from input directory
                rel_path = md_file.relative_to(input_dir)

                # Change extension to .docx
                docx_name = rel_path.with_suffix(".docx")
                output_path = self.output_dir / docx_name

                future = executor.submit(self.convert_file, md_file, output_path)
                future_to_file[future] = md_file

            # Process completed tasks
            for future in as_completed(future_to_file):
                md_file = future_to_file[future]
                try:
                    success = future.result()
                    if success:
                        successful += 1
                    else:
                        failed += 1
                except Exception as e:
                    self.logger.error(f"Unexpected error processing {md_file}: {e}")
                    failed += 1

        return {"total": len(md_files), "successful": successful, "failed": failed}

    def save_metadata(self) -> None:
        """Save conversion metadata to files."""
        # Ensure metadata directory exists
        METADATA_DIR.mkdir(parents=True, exist_ok=True)

        # Save conversion log
        log_file = METADATA_DIR / "conversion-log.json"
        with open(log_file, "w", encoding="utf-8") as f:
            json.dump(self.conversion_log, f, indent=2, ensure_ascii=False)

        # Save failed conversions
        if self.failed_conversions:
            failed_file = METADATA_DIR / "failed-conversions.json"
            with open(failed_file, "w", encoding="utf-8") as f:
                json.dump(self.failed_conversions, f, indent=2, ensure_ascii=False)

        # Create document index
        index = {
            "conversion_date": datetime.now().isoformat(),
            "total_files": len(self.conversion_log),
            "successful_conversions": len(self.conversion_log),
            "failed_conversions": len(self.failed_conversions),
            "phases": {},
        }

        # Group by phase
        for log_entry in self.conversion_log:
            input_path = Path(log_entry["input"])
            phase = self.get_phase_from_path(input_path)
            if phase:
                if phase not in index["phases"]:
                    index["phases"][phase] = []
                index["phases"][phase].append(
                    {
                        "title": log_entry["metadata"]["title"],
                        "input": log_entry["input"],
                        "output": log_entry["output"],
                    }
                )

        index_file = METADATA_DIR / "document-index.json"
        with open(index_file, "w", encoding="utf-8") as f:
            json.dump(index, f, indent=2, ensure_ascii=False)


def main():
    """Main function."""
    parser = argparse.ArgumentParser(
        description="Convert [CLIENT] F&I documentation to DOCX"
    )
    parser.add_argument(
        "--phase", help="Process only specific phase (e.g., 9-contract)"
    )
    parser.add_argument("--file", help="Process only specific file")
    parser.add_argument(
        "--parallel", type=int, default=4, help="Number of parallel conversions"
    )
    parser.add_argument("--verbose", "-v", action="store_true", help="Verbose output")
    parser.add_argument("--force", action="store_true", help="Overwrite existing files")

    args = parser.parse_args()

    # Setup paths
    base_dir = Path.cwd()
    opportunity_dir = base_dir / OPPORTUNITY_DIR
    output_dir = base_dir / OUTPUT_DIR

    # Check if opportunity directory exists
    if not opportunity_dir.exists():
        print(f"Error: {opportunity_dir} directory not found")
        sys.exit(1)

    # Create processor
    processor = DocumentProcessor(output_dir, args.verbose)

    print("Converting [CLIENT] F&I documentation to DOCX...")
    print(f"Parallel workers: {args.parallel}")

    # Process specific file, phase, or all
    if args.file:
        file_path = Path(args.file)
        if not file_path.exists():
            print(f"Error: File not found: {file_path}")
            sys.exit(1)

        print(f"Processing file: {file_path}")
        # Calculate output path
        if file_path.is_relative_to(opportunity_dir):
            rel_path = file_path.relative_to(opportunity_dir)
        else:
            rel_path = Path(file_path.name)
        output_path = output_dir / rel_path.with_suffix(".docx")

        success = processor.convert_file(file_path, output_path)
        results = {
            "total": 1,
            "successful": 1 if success else 0,
            "failed": 0 if success else 1,
        }
    elif args.phase:
        phase_dir = opportunity_dir / args.phase
        if not phase_dir.exists():
            print(f"Error: Phase directory not found: {phase_dir}")
            sys.exit(1)

        print(f"Processing phase: {args.phase}")
        results = processor.process_directory(phase_dir, args.parallel)
    else:
        print("Processing all phases...")
        results = processor.process_directory(opportunity_dir, args.parallel)

    # Save metadata
    processor.save_metadata()

    # Print results
    print("\nConversion Results:")
    print(f"Total files: {results['total']}")
    print(f"Successful: {results['successful']}")
    print(f"Failed: {results['failed']}")

    if results["failed"] > 0:
        print(
            f"\nFailed conversions logged to: {METADATA_DIR / 'failed-conversions.json'}"
        )

    print(f"\nOutput directory: {output_dir}")
    print(f"Metadata saved to: {METADATA_DIR}")


if __name__ == "__main__":
    main()
