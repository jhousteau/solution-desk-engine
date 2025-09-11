#!/usr/bin/env python3
"""
Branding Manager for Client Documentation
Centralized management of Capgemini brand assets, colors, and styling guidelines.
Ensures consistent corporate branding across all document formats.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional, Tuple

try:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    pass


class BrandingManager:
    """Centralized branding and style management for all document formats."""

    def __init__(self, branding_dir: Path):
        self.branding_dir = branding_dir
        self.logger = logging.getLogger(__name__)

        # Brand asset paths
        self.logo_kit_dir = (
            branding_dir / "Capgemini_Logo_Basic-Kit" / "Capgemini_Logo_kit"
        )
        self.logo_color_png = self.logo_kit_dir / "Capgemini_Logo_Color.png"
        self.logo_color_jpg = self.logo_kit_dir / "Capgemini_Logo_Color_Print.jpg"
        self.pptx_template = branding_dir / "Capgemini-template-master_2024-2.potx"

        # Verify assets exist
        self._verify_brand_assets()

    def _verify_brand_assets(self) -> None:
        """Verify that required brand assets are available."""
        missing_assets = []

        if not self.logo_color_png.exists():
            missing_assets.append(str(self.logo_color_png))

        if missing_assets:
            self.logger.warning(f"Missing brand assets: {missing_assets}")

    @property
    def brand_colors_rgb(self) -> Dict[str, RGBColor]:
        """Capgemini brand colors as RGBColor objects for DOCX."""
        return {
            "primary_blue": RGBColor(0, 112, 173),  # #0070AD - Capgemini Primary Blue
            "secondary_teal": RGBColor(
                18, 171, 219
            ),  # #12ABDB - Capgemini Secondary Teal
            "dark_blue": RGBColor(0, 51, 161),  # #0033A1 - Capgemini Dark Blue
            "text_gray": RGBColor(51, 51, 51),  # #333333 - Standard Text Gray
            "light_gray": RGBColor(245, 245, 245),  # #F5F5F5 - Light Background Gray
            "medium_gray": RGBColor(128, 128, 128),  # #808080 - Medium Gray
            "white": RGBColor(255, 255, 255),  # #FFFFFF - White
            "accent_orange": RGBColor(255, 102, 0),  # #FF6600 - Accent Orange
            "success_green": RGBColor(76, 175, 80),  # #4CAF50 - Success Green
            "warning_amber": RGBColor(255, 193, 7),  # #FFC107 - Warning Amber
        }

    @property
    def brand_colors_hex(self) -> Dict[str, str]:
        """Capgemini brand colors as hex codes for CSS/PDF."""
        return {
            "primary_blue": "#0070AD",
            "secondary_teal": "#12ABDB",
            "dark_blue": "#0033A1",
            "text_gray": "#333333",
            "light_gray": "#F5F5F5",
            "medium_gray": "#808080",
            "white": "#FFFFFF",
            "accent_orange": "#FF6600",
            "success_green": "#4CAF50",
            "warning_amber": "#FFC107",
        }

    @property
    def brand_fonts(self) -> Dict[str, str]:
        """Capgemini brand typography specifications."""
        return {
            "primary": "Arial",  # Primary corporate font
            "secondary": "Helvetica",  # Secondary font
            "monospace": "Courier New",  # Code/technical content
            "fallback": "sans-serif",  # Web fallback
        }

    @property
    def typography_scale(self) -> Dict[str, int]:
        """Typography scale for consistent sizing."""
        return {
            "h1": 24,  # Main titles
            "h2": 20,  # Section headers
            "h3": 16,  # Subsection headers
            "h4": 14,  # Minor headers
            "body": 11,  # Body text
            "small": 9,  # Footer/caption text
            "code": 10,  # Code text
        }

    def get_phase_color(
        self, phase_key: str, format_type: str = "rgb"
    ) -> Optional[object]:
        """Get phase-specific brand color."""
        phase_colors = {
            "0-source": "dark_blue",
            "1-research": "primary_blue",
            "2-requirements": "primary_blue",
            "3-analysis": "secondary_teal",
            "4-business-case": "dark_blue",
            "5-architecture": "primary_blue",
            "6-solution-design": "secondary_teal",
            "7-implementation-plan": "dark_blue",
            "8-proposal": "primary_blue",
            "9-contract": "dark_blue",
            "10-audit": "text_gray",
        }

        color_key = phase_colors.get(phase_key, "primary_blue")

        if format_type == "rgb":
            return self.brand_colors_rgb.get(color_key)
        elif format_type == "hex":
            return self.brand_colors_hex.get(color_key)
        else:
            return color_key

    def apply_docx_header_footer(
        self, doc, project_title: str = "Client Documentation"
    ) -> None:
        """Apply branded header and footer to DOCX document."""
        try:
            section = doc.sections[0]

            # Header with logo
            header = section.header
            if self.logo_color_png.exists():
                header_para = header.paragraphs[0]
                run = header_para.runs[0] if header_para.runs else header_para.add_run()
                run.add_picture(str(self.logo_color_png), width=Inches(1.5))
                header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Footer with branding
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Confidential - Capgemini & [CLIENT] | {project_title}"
            footer_para.style.font.size = Pt(self.typography_scale["small"])
            footer_para.style.font.color.rgb = self.brand_colors_rgb["text_gray"]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        except Exception as e:
            self.logger.warning(f"Could not apply header/footer: {e}")

    def setup_docx_styles(self, doc) -> None:
        """Setup branded styles for DOCX document."""
        try:
            # Normal style
            normal_style = doc.styles["Normal"]
            normal_style.font.name = self.brand_fonts["primary"]
            normal_style.font.size = Pt(self.typography_scale["body"])
            normal_style.font.color.rgb = self.brand_colors_rgb["text_gray"]

            # Heading styles with brand colors
            headings = [
                ("Heading 1", "h1", "primary_blue"),
                ("Heading 2", "h2", "dark_blue"),
                ("Heading 3", "h3", "text_gray"),
                ("Heading 4", "h4", "text_gray"),
            ]

            for style_name, size_key, color_key in headings:
                if style_name in doc.styles:
                    heading_style = doc.styles[style_name]
                    heading_style.font.name = self.brand_fonts["primary"]
                    heading_style.font.size = Pt(self.typography_scale[size_key])
                    heading_style.font.color.rgb = self.brand_colors_rgb[color_key]
                    heading_style.font.bold = True

        except Exception as e:
            self.logger.warning(f"Could not setup DOCX styles: {e}")

    def generate_pdf_css(self, phase: str = None) -> str:
        """Generate branded CSS for PDF conversion."""
        phase_color = (
            self.get_phase_color(phase, "hex")
            if phase
            else self.brand_colors_hex["primary_blue"]
        )

        return f"""
/* Capgemini Brand Styles for PDF */
@page {{
    margin: 1in;
    @bottom-center {{
        content: "Confidential - Capgemini & [CLIENT]";
        font-size: {self.typography_scale["small"]}pt;
        color: {self.brand_colors_hex["text_gray"]};
    }}
}}

body {{
    font-family: "{self.brand_fonts["primary"]}", {self.brand_fonts["fallback"]};
    font-size: {self.typography_scale["body"]}pt;
    color: {self.brand_colors_hex["text_gray"]};
    line-height: 1.6;
    margin: 0;
    padding: 0;
}}

/* Headings with brand colors */
h1 {{
    color: {self.brand_colors_hex["primary_blue"]};
    font-size: {self.typography_scale["h1"]}pt;
    font-weight: bold;
    margin-top: 24pt;
    margin-bottom: 12pt;
    page-break-after: avoid;
}}

h2 {{
    color: {self.brand_colors_hex["dark_blue"]};
    font-size: {self.typography_scale["h2"]}pt;
    font-weight: bold;
    margin-top: 18pt;
    margin-bottom: 9pt;
    page-break-after: avoid;
}}

h3 {{
    color: {self.brand_colors_hex["text_gray"]};
    font-size: {self.typography_scale["h3"]}pt;
    font-weight: bold;
    margin-top: 14pt;
    margin-bottom: 7pt;
}}

h4 {{
    color: {self.brand_colors_hex["text_gray"]};
    font-size: {self.typography_scale["h4"]}pt;
    font-weight: bold;
    margin-top: 12pt;
    margin-bottom: 6pt;
}}

/* Phase-specific accent color */
.phase-header {{
    border-left: 4px solid {phase_color};
    padding-left: 12pt;
    margin-left: 0;
}}

/* Tables with brand styling */
table {{
    border-collapse: collapse;
    width: 100%;
    margin: 12pt 0;
}}

th {{
    background-color: {self.brand_colors_hex["primary_blue"]};
    color: {self.brand_colors_hex["white"]};
    font-weight: bold;
    padding: 8pt;
    text-align: left;
    border: 1pt solid {self.brand_colors_hex["medium_gray"]};
}}

td {{
    padding: 6pt 8pt;
    border: 1pt solid {self.brand_colors_hex["medium_gray"]};
}}

tr:nth-child(even) {{
    background-color: {self.brand_colors_hex["light_gray"]};
}}

/* Code blocks */
pre, code {{
    font-family: "{self.brand_fonts["monospace"]}", monospace;
    font-size: {self.typography_scale["code"]}pt;
    background-color: {self.brand_colors_hex["light_gray"]};
    border-left: 3px solid {self.brand_colors_hex["secondary_teal"]};
}}

pre {{
    padding: 8pt;
    margin: 8pt 0;
    overflow-x: auto;
}}

code {{
    padding: 2pt 4pt;
}}

/* Blockquotes */
blockquote {{
    border-left: 4px solid {self.brand_colors_hex["secondary_teal"]};
    margin: 12pt 0;
    padding-left: 12pt;
    font-style: italic;
    color: {self.brand_colors_hex["text_gray"]};
}}

/* Lists */
ul, ol {{
    margin: 8pt 0;
    padding-left: 24pt;
}}

li {{
    margin: 4pt 0;
}}

/* Links */
a {{
    color: {self.brand_colors_hex["primary_blue"]};
    text-decoration: underline;
}}

a:hover {{
    color: {self.brand_colors_hex["dark_blue"]};
}}

/* Page breaks */
.page-break {{
    page-break-before: always;
}}

/* Brand accent elements */
.brand-accent {{
    border-left: 4px solid {self.brand_colors_hex["primary_blue"]};
    background-color: {self.brand_colors_hex["light_gray"]};
    padding: 8pt 12pt;
    margin: 12pt 0;
}}

.highlight {{
    background-color: {self.brand_colors_hex["warning_amber"]};
    padding: 2pt 4pt;
}}

.success {{
    color: {self.brand_colors_hex["success_green"]};
    font-weight: bold;
}}

.warning {{
    color: {self.brand_colors_hex["accent_orange"]};
    font-weight: bold;
}}
"""

    def create_brand_cover_page_content(
        self, title: str, subtitle: str = None, date: str = None
    ) -> str:
        """Generate branded cover page content for any format."""
        current_date = date or datetime.now().strftime("%Y-%m-%d")
        subtitle = subtitle or "Finance & Insurance AI Assistant"

        return f"""# Statement of Work

## [CLIENT]
### {subtitle}

**{title}**

---

*Google Cloud Deal Acceleration Fund (DAF) Implementation*

**Date:** {current_date}
**Version:** 1.0
**Prepared by:** Capgemini

---

**Confidential Document**
This document contains confidential and proprietary information of Capgemini and [CLIENT] Automotive Group.
"""

    def get_brand_metadata(self, title: str, phase: str = None) -> Dict[str, str]:
        """Generate branded document metadata."""
        return {
            "title": title,
            "author": "Capgemini",
            "company": "Capgemini & [CLIENT] Automotive Group",
            "subject": "F&I AI Assistant - Proof of Value Implementation",
            "keywords": "AI, Finance, Insurance, Automotive, [CLIENT], Capgemini, DAF",
            "creator": "Capgemini Document Conversion System",
            "date": datetime.now().strftime("%Y-%m-%d"),
            "version": "1.0",
            "confidentiality": "Confidential",
            "phase": phase or "documentation",
            "project": "[CLIENT] F&I AI Assistant POV",
        }

    def create_google_drive_folder_structure(self) -> Dict[str, str]:
        """Define branded Google Drive folder organization."""
        return {
            "root": "[CLIENT] F&I Documentation - Capgemini DAF Implementation",
            "phases": {
                "0-source": "📋 Source Materials",
                "1-research": "🔍 Research & Discovery",
                "2-requirements": "📝 Requirements Analysis",
                "3-analysis": "👥 Stakeholder Analysis",
                "4-business-case": "💼 Business Case",
                "5-architecture": "🏗️ Architecture Design",
                "6-solution-design": "⚙️ Solution Design",
                "7-implementation-plan": "📋 Implementation Plan",
                "8-proposal": "📊 Proposal Materials",
                "9-contract": "📄 Contract Documents",
                "10-audit": "✅ Audit & Validation",
            },
            "metadata": "📊 Project Metadata",
        }

    def apply_brand_validation(self, content: str) -> Tuple[bool, list]:
        """Validate content meets brand guidelines."""
        violations = []

        # Check for consistent terminology
        brand_terms = {
            "AI Assistant": ["ai assistant", "AI assistant"],
            "Proof of Value": ["proof of value", "POV", "proof-of-value"],
            "Finance & Insurance": [
                "F&I",
                "finance and insurance",
                "Finance and Insurance",
            ],
        }

        # Validate brand term consistency
        for preferred, alternatives in brand_terms.items():
            for alt in alternatives:
                if alt in content.lower() and preferred not in content:
                    violations.append(f"Use '{preferred}' instead of '{alt}'")

        # Check for required confidentiality notice
        if "confidential" not in content.lower():
            violations.append("Missing confidentiality notice")

        is_valid = len(violations) == 0
        return is_valid, violations
